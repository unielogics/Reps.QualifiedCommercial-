from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "FIELD_DESK_APPLICATION_TRAINING.md"
OUTPUT = ROOT / "docs" / "training-manual" / "Field_Desk_Application_Workflow_Training_Booklet.docx"

NAVY = "17365D"
BLUE = "2454A6"
LIGHT_BLUE = "EAF1FC"
MUTED = "657184"
INK = "111924"
LINE = "DCE2EA"
GREEN = "147A53"
LIGHT_GREEN = "E9F6EF"
AMBER = "8A5A00"
LIGHT_AMBER = "FFF4DD"


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill, border_color=None):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr")
            ppr.append(pbdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        pbdr.append(left)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8, color=MUTED)


def add_bottom_rule(paragraph, color=LINE, size="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name, left, hanging in (
        ("List Bullet", 0.375, 0.188),
        ("List Number", 0.375, 0.188),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(-hanging)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Figure Caption" not in [style.name for style in doc.styles]:
        caption = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(4)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = hp.add_run("QUALIFIED COMMERCIAL  |  FIELD DESK")
    set_run_font(left, size=8, color=MUTED, bold=True)
    right = hp.add_run("\tAPPLICATION WORKFLOW")
    set_run_font(right, size=8, color=MUTED, bold=True)
    add_bottom_rule(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    lead = fp.add_run("Internal training  |  Sanitized examples  |  ")
    set_run_font(lead, size=8, color=MUTED)
    add_page_field(fp, "PAGE")
    mid = fp.add_run(" of ")
    set_run_font(mid, size=8, color=MUTED)
    add_page_field(fp, "NUMPAGES")


def add_inline_markup(paragraph, text, size=11, color=INK):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_callout(doc, label, body, tone="blue"):
    colors = {
        "blue": (LIGHT_BLUE, BLUE),
        "green": (LIGHT_GREEN, GREEN),
        "amber": (LIGHT_AMBER, AMBER),
    }
    fill, border = colors[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:color"), border)
    left.set(qn("w:space"), "6")
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label} ")
    set_run_font(r, size=10.5, color=border, bold=True)
    add_inline_markup(p, body, size=10.5, color=INK)
    return p


def add_picture(doc, image_path, alt_text, figure_number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(5.75))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(f"Figure {figure_number}. {alt_text}")


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("EMPLOYEE TRAINING BOOKLET")
    set_run_font(r, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Field Desk Application Workflow")
    set_run_font(r, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    r = subtitle.add_run("From first business visit through signed application and super-admin handoff")
    set_run_font(r, size=14, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    add_bottom_rule(rule, color=BLUE, size="18")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(64)
    r = meta.add_run("Version: August 2026   |   Audience: Field representatives, loan executives, and super admins")
    set_run_font(r, size=10.5, color=MUTED, bold=True)

    add_callout(
        doc,
        "Training rule:",
        "Follow the application gates in order. Correct missing information instead of bypassing a step. The client signs independently; staff never apply a client's signature.",
        tone="amber",
    )

    privacy = doc.add_paragraph()
    privacy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    privacy.paragraph_format.space_before = Pt(48)
    r = privacy.add_run("All names, addresses, amounts, and account details shown in this guide are sanitized training examples.")
    set_run_font(r, size=9, color=MUTED, italic=True)
    doc.add_page_break()


def parse_manual(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    figure_number = 0
    started = False
    first_numbered_section = True

    for raw in lines:
        line = raw.strip()
        if not started:
            if line == "## Workflow at a Glance":
                started = True
            else:
                continue

        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            if re.match(r"\d+\.", title) and first_numbered_section:
                doc.add_page_break()
                first_numbered_section = False
            p = doc.add_paragraph(title, style="Heading 1")
            p.paragraph_format.keep_with_next = True
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        image_match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            figure_number += 1
            alt_text, relative = image_match.groups()
            add_picture(doc, SOURCE.parent / relative, alt_text, figure_number)
            continue
        ordered = re.match(r"(\d+)\.\s+(.+)", line)
        if ordered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            marker = p.add_run(f"{ordered.group(1)}. ")
            set_run_font(marker, size=10.5, color=INK)
            add_inline_markup(p, ordered.group(2), size=10.5)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, line[2:])
            continue
        if line.startswith("**") and ":**" in line:
            match = re.match(r"\*\*(.+?):\*\*\s*(.*)", line)
            if match:
                label, body = match.groups()
                if label == "Purpose":
                    add_callout(doc, "Purpose:", body, tone="blue")
                elif label == "Outcome":
                    add_callout(doc, "Outcome:", body, tone="green")
                else:
                    p = doc.add_paragraph(style="Heading 2")
                    p.add_run(label)
                    if body:
                        detail = doc.add_paragraph()
                        add_inline_markup(detail, body)
                continue
        if line.startswith("**") and line.endswith("**"):
            doc.add_paragraph(line[2:-2], style="Heading 2")
            continue
        p = doc.add_paragraph()
        add_inline_markup(p, line)


def add_closing_page(doc):
    doc.add_page_break()
    doc.add_paragraph("Final Agent Checklist", style="Heading 1")
    checks = [
        "The application carries the exact six-digit NAICS activity selected during intake.",
        "Ownership totals exactly 100%, and every 20%+ owner has unique personal contact information.",
        "Every required owner completed an independent soft-credit authorization and pull.",
        "At least three current months of official bank-produced statements are verified.",
        "Step 3 metrics are supported by evidence; unavailable values are not guessed.",
        "Three client review windows are saved as preferences only.",
        "The Step 4 underwriting package is complete for the viable route.",
        "The exact populated QC application was reviewed with the client before sending.",
        "The client signed independently, and the executed PDF/certificate is visible.",
        "The file is ready for super-admin decision, calendar invitation, bucket follow-up, and handoff.",
    ]
    for item in checks:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markup(p, item)

    add_callout(
        doc,
        "Escalate instead of guessing:",
        "Use the file's Desk conversation for internal clarification. Do not alter eligibility answers, evidence classification, signature records, or client contact details to force a gate open.",
        tone="amber",
    )


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    parse_manual(doc)
    add_closing_page(doc)
    doc.core_properties.title = "Field Desk Application Workflow Training Booklet"
    doc.core_properties.subject = "Employee training for the Qualified Commercial Field Desk application workflow"
    doc.core_properties.author = "Qualified Commercial"
    doc.core_properties.keywords = "Field Desk, application, iSoftPull, banking, underwriting, signature, training"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
